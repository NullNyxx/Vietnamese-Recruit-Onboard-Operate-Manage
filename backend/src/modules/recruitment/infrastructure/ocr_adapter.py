"""OCR Adapter for olmOCR integration.

Communicates with the olmOCR server (Vision Language Model) for text extraction
from PDF, DOCX, and image files. Handles PDF chunking for large documents,
DOCX text extraction with fallback to OCR, and retry logic with exponential backoff.
"""

from __future__ import annotations

import asyncio
import io
import logging
from typing import Final

import httpx

from src.modules.recruitment.domain.exceptions import OCRExtractionError
from src.modules.recruitment.infrastructure.config import RecruitmentSettings

logger = logging.getLogger(__name__)

# MIME type constants
_PDF_MIME: Final[str] = "application/pdf"
_DOCX_MIME: Final[str] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_IMAGE_MIMES: Final[set[str]] = {"image/jpeg", "image/png"}

# Separator used between PDF chunks after OCR
_CHUNK_SEPARATOR: Final[str] = "\n---PAGE BREAK---\n"

# Threshold for DOCX fallback to OCR
_DOCX_MIN_TEXT_LENGTH: Final[int] = 50
_DOCX_MIN_FILE_SIZE: Final[int] = 500 * 1024  # 500KB

# Base backoff delay in seconds
_BASE_BACKOFF_SECONDS: Final[float] = 5.0


class OCRAdapter:
    """Communicates with olmOCR server for text extraction.

    Supports PDF (with chunking for large files), DOCX (with fallback to
    PDF conversion + OCR), and image files (JPEG/PNG).

    Attributes:
        _settings: Recruitment module configuration.
    """

    def __init__(self, settings: RecruitmentSettings) -> None:
        """Initialize OCRAdapter with recruitment settings.

        Args:
            settings: RecruitmentSettings instance with olmOCR configuration.
        """
        self._settings = settings

    async def extract_text(self, file_content: bytes, filename: str, mime_type: str) -> str:
        """Extract text from a file using olmOCR.

        Routes to the appropriate handler based on MIME type:
        - PDF: direct send or chunked processing for large files
        - DOCX: python-docx extraction with OCR fallback
        - Image (JPEG/PNG): direct send to olmOCR

        Args:
            file_content: Raw file bytes.
            filename: Original filename (used for multipart upload).
            mime_type: MIME type of the file.

        Returns:
            Extracted markdown text from the file.

        Raises:
            OCRExtractionError: If text extraction fails after all retries.
        """
        try:
            if mime_type == _PDF_MIME:
                return await self._handle_pdf(file_content, filename)
            elif mime_type == _DOCX_MIME:
                return await self._handle_docx(file_content, filename)
            elif mime_type in _IMAGE_MIMES:
                return await self._send_to_ocr(file_content, filename, mime_type)
            else:
                raise OCRExtractionError(f"Unsupported MIME type for OCR: {mime_type}")
        except OCRExtractionError:
            raise
        except Exception as exc:
            logger.error(
                "Unexpected error during OCR extraction for file '%s': %s",
                filename,
                exc,
            )
            raise OCRExtractionError(f"OCR extraction failed for '{filename}': {exc}") from exc

    async def _handle_pdf(self, file_content: bytes, filename: str) -> str:
        """Handle PDF files, splitting into chunks if necessary.

        Args:
            file_content: Raw PDF bytes.
            filename: Original filename.

        Returns:
            Extracted markdown text (concatenated chunks if split).

        Raises:
            OCRExtractionError: If OCR fails after retries.
        """
        import fitz  # PyMuPDF

        max_pages = self._settings.olmocr_max_pages_per_chunk

        doc = fitz.open(stream=file_content, filetype="pdf")
        try:
            page_count = doc.page_count

            if page_count <= max_pages:
                # Send entire PDF directly
                return await self._send_to_ocr(file_content, filename, _PDF_MIME)

            # Split into chunks and process each
            logger.info(
                "PDF '%s' has %d pages, splitting into chunks of %d",
                filename,
                page_count,
                max_pages,
            )

            chunks_text: list[str] = []
            for chunk_start in range(0, page_count, max_pages):
                chunk_end = min(chunk_start + max_pages, page_count)

                # Create a new PDF with just this chunk's pages
                chunk_doc = fitz.open()
                try:
                    chunk_doc.insert_pdf(doc, from_page=chunk_start, to_page=chunk_end - 1)
                    chunk_bytes = chunk_doc.tobytes()
                finally:
                    chunk_doc.close()

                chunk_filename = f"{filename}_chunk_{chunk_start + 1}-{chunk_end}.pdf"
                chunk_text = await self._send_to_ocr(chunk_bytes, chunk_filename, _PDF_MIME)
                chunks_text.append(chunk_text)

            return _CHUNK_SEPARATOR.join(chunks_text)
        finally:
            doc.close()

    async def _handle_docx(self, file_content: bytes, filename: str) -> str:
        """Handle DOCX files with python-docx extraction and OCR fallback.

        First attempts to extract text directly using python-docx. If the
        extracted text is less than 50 characters AND the file is larger than
        500KB, falls back to converting to PDF and sending to OCR.

        Args:
            file_content: Raw DOCX bytes.
            filename: Original filename.

        Returns:
            Extracted text (either from python-docx or OCR).

        Raises:
            OCRExtractionError: If both extraction methods fail.
        """
        from docx import Document

        # Try direct text extraction with python-docx
        try:
            doc = Document(io.BytesIO(file_content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            extracted_text = "\n".join(paragraphs)
        except Exception as exc:
            logger.warning(
                "python-docx extraction failed for '%s': %s, falling back to OCR",
                filename,
                exc,
            )
            extracted_text = ""

        # Check if fallback to OCR is needed
        file_size = len(file_content)
        if len(extracted_text) >= _DOCX_MIN_TEXT_LENGTH:
            return extracted_text

        if file_size <= _DOCX_MIN_FILE_SIZE:
            # File is small and text is short — return what we have
            return extracted_text

        # Fallback: convert DOCX to PDF and OCR
        logger.info(
            "DOCX '%s' has insufficient text (%d chars) and is large (%d bytes), "
            "converting to PDF for OCR",
            filename,
            len(extracted_text),
            file_size,
        )
        pdf_bytes = await self._convert_docx_to_pdf(file_content, filename)
        return await self._send_to_ocr(pdf_bytes, filename.rsplit(".", 1)[0] + ".pdf", _PDF_MIME)

    async def _convert_docx_to_pdf(self, file_content: bytes, filename: str) -> bytes:
        """Convert DOCX to PDF for OCR processing.

        Uses PyMuPDF's built-in story/document capabilities or falls back
        to a simple approach using the DOCX text content.

        Args:
            file_content: Raw DOCX bytes.
            filename: Original filename for logging.

        Returns:
            PDF bytes suitable for OCR.

        Raises:
            OCRExtractionError: If conversion fails.
        """
        try:
            import fitz  # PyMuPDF
            from docx import Document

            # Extract all text from DOCX
            doc = Document(io.BytesIO(file_content))
            full_text = "\n".join(p.text for p in doc.paragraphs)

            # Create a simple PDF with the text content using PyMuPDF
            pdf_doc = fitz.open()
            try:
                # Create pages with text content
                # Use A4 page size
                rect = fitz.Rect(0, 0, 595, 842)
                page = pdf_doc.new_page(width=rect.width, height=rect.height)

                # Insert text with wrapping
                text_rect = fitz.Rect(50, 50, 545, 792)
                page.insert_textbox(
                    text_rect,
                    full_text,
                    fontsize=10,
                    fontname="helv",
                )

                pdf_bytes = pdf_doc.tobytes()
            finally:
                pdf_doc.close()

            return pdf_bytes
        except Exception as exc:
            logger.error("Failed to convert DOCX '%s' to PDF: %s", filename, exc)
            raise OCRExtractionError(
                f"DOCX to PDF conversion failed for '{filename}': {exc}"
            ) from exc

    async def _send_to_ocr(self, file_content: bytes, filename: str, mime_type: str) -> str:
        """Send file to olmOCR server with retry logic.

        Posts the file as multipart/form-data to the olmOCR endpoint.
        Retries up to max_retries times with exponential backoff (5s, 10s, 20s).

        Args:
            file_content: Raw file bytes to send.
            filename: Filename for the multipart upload.
            mime_type: MIME type of the file.

        Returns:
            Extracted markdown text from olmOCR response.

        Raises:
            OCRExtractionError: If all retry attempts fail.
        """
        endpoint = self._settings.olmocr_endpoint_url
        timeout = self._settings.olmocr_timeout_seconds
        max_retries = self._settings.olmocr_max_retries

        last_error: Exception | None = None

        for attempt in range(max_retries):
            try:
                async with httpx.AsyncClient(
                    timeout=httpx.Timeout(timeout, connect=30.0)
                ) as client:
                    files = {
                        "file": (filename, file_content, mime_type),
                    }
                    response = await client.post(endpoint, files=files)

                if response.status_code != 200:
                    error_msg = f"olmOCR returned HTTP {response.status_code} for file '{filename}'"
                    logger.warning("%s (attempt %d/%d)", error_msg, attempt + 1, max_retries)
                    last_error = OCRExtractionError(error_msg)
                else:
                    # Parse response JSON
                    data = response.json()
                    markdown = data.get("markdown")
                    if markdown is None:
                        error_msg = (
                            f"olmOCR response missing 'markdown' field for file '{filename}'"
                        )
                        logger.warning(
                            "%s (attempt %d/%d)",
                            error_msg,
                            attempt + 1,
                            max_retries,
                        )
                        last_error = OCRExtractionError(error_msg)
                    else:
                        return markdown

            except httpx.ConnectError as exc:
                error_msg = f"olmOCR server unreachable for file '{filename}': {exc}"
                logger.warning("%s (attempt %d/%d)", error_msg, attempt + 1, max_retries)
                last_error = exc

            except httpx.TimeoutException as exc:
                error_msg = f"olmOCR request timed out for file '{filename}': {exc}"
                logger.warning("%s (attempt %d/%d)", error_msg, attempt + 1, max_retries)
                last_error = exc

            except Exception as exc:
                error_msg = f"Unexpected error calling olmOCR for file '{filename}': {exc}"
                logger.warning("%s (attempt %d/%d)", error_msg, attempt + 1, max_retries)
                last_error = exc

            # Exponential backoff: 5s * 2^attempt → 5s, 10s, 20s
            if attempt < max_retries - 1:
                backoff = _BASE_BACKOFF_SECONDS * (2**attempt)
                logger.info(
                    "Retrying olmOCR for '%s' in %.1fs (attempt %d/%d)",
                    filename,
                    backoff,
                    attempt + 2,
                    max_retries,
                )
                await asyncio.sleep(backoff)

        # All retries exhausted
        raise OCRExtractionError(
            f"OCR extraction failed for '{filename}' after {max_retries} attempts: {last_error}"
        )
