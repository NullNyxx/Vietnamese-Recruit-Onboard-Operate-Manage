"""Tests for OCR Adapter (olmOCR integration)."""

from unittest.mock import AsyncMock, MagicMock, patch

import httpx
import pytest

from src.modules.recruitment.domain.exceptions import OCRExtractionError
from src.modules.recruitment.infrastructure.config import RecruitmentSettings
from src.modules.recruitment.infrastructure.ocr_adapter import OCRAdapter


@pytest.fixture
def settings(monkeypatch: pytest.MonkeyPatch) -> RecruitmentSettings:
    monkeypatch.setenv("RECRUITMENT_OLMOCR_ENDPOINT_URL", "https://olmocr.test/ocr")
    monkeypatch.setenv("RECRUITMENT_OLMOCR_TIMEOUT_SECONDS", "600")
    monkeypatch.setenv("RECRUITMENT_OLMOCR_MAX_RETRIES", "3")
    monkeypatch.setenv("RECRUITMENT_OLMOCR_MAX_PAGES_PER_CHUNK", "20")
    return RecruitmentSettings()


@pytest.fixture
def adapter(settings: RecruitmentSettings) -> OCRAdapter:
    return OCRAdapter(settings)


def _make_response(status_code: int = 200, json_data: dict | None = None) -> httpx.Response:
    """Helper to create a mock httpx.Response."""
    response = httpx.Response(
        status_code=status_code,
        json=json_data if json_data else {},
        request=httpx.Request("POST", "https://olmocr.test/ocr"),
    )
    return response


class TestExtractTextRouting:
    """Verify extract_text routes to correct handler based on MIME type."""

    async def test_pdf_routes_to_handle_pdf(self, adapter: OCRAdapter) -> None:
        with patch.object(adapter, "_handle_pdf", new_callable=AsyncMock) as mock_pdf:
            mock_pdf.return_value = "pdf text"
            result = await adapter.extract_text(b"pdf-data", "test.pdf", "application/pdf")
        assert result == "pdf text"
        mock_pdf.assert_called_once_with(b"pdf-data", "test.pdf")

    async def test_docx_routes_to_handle_docx(self, adapter: OCRAdapter) -> None:
        docx_mime = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        with patch.object(adapter, "_handle_docx", new_callable=AsyncMock) as mock_docx:
            mock_docx.return_value = "docx text"
            result = await adapter.extract_text(b"docx-data", "test.docx", docx_mime)
        assert result == "docx text"
        mock_docx.assert_called_once_with(b"docx-data", "test.docx")

    async def test_jpeg_routes_to_send_to_ocr(self, adapter: OCRAdapter) -> None:
        with patch.object(adapter, "_send_to_ocr", new_callable=AsyncMock) as mock_ocr:
            mock_ocr.return_value = "image text"
            result = await adapter.extract_text(b"jpeg-data", "photo.jpg", "image/jpeg")
        assert result == "image text"
        mock_ocr.assert_called_once_with(b"jpeg-data", "photo.jpg", "image/jpeg")

    async def test_png_routes_to_send_to_ocr(self, adapter: OCRAdapter) -> None:
        with patch.object(adapter, "_send_to_ocr", new_callable=AsyncMock) as mock_ocr:
            mock_ocr.return_value = "png text"
            result = await adapter.extract_text(b"png-data", "scan.png", "image/png")
        assert result == "png text"
        mock_ocr.assert_called_once_with(b"png-data", "scan.png", "image/png")

    async def test_unsupported_mime_raises_error(self, adapter: OCRAdapter) -> None:
        with pytest.raises(OCRExtractionError, match="Unsupported MIME type"):
            await adapter.extract_text(b"data", "file.txt", "text/plain")


class TestSendToOCR:
    """Verify _send_to_ocr HTTP interaction and retry logic."""

    async def test_successful_response_returns_markdown(self, adapter: OCRAdapter) -> None:
        mock_response = _make_response(200, {"markdown": "# Hello\nExtracted text"})

        with patch("httpx.AsyncClient") as mock_client_cls:
            mock_client = AsyncMock()
            mock_client.post = AsyncMock(return_value=mock_response)
            mock_client_cls.return_value.__aenter__ = AsyncMock(return_value=mock_client)
            mock_client_cls.return_value.__aexit__ = AsyncMock(return_value=False)

            result = await adapter._send_to_ocr(b"file-data", "test.pdf", "application/pdf")

        assert result == "# Hello\nExtracted text"

    async def test_non_200_retries_and_raises(self, adapter: OCRAdapter) -> None:
        mock_response = _make_response(500, {"error": "internal"})

        with patch("httpx.AsyncClient") as mock_client_cls:
            mock_client = AsyncMock()
            mock_client.post = AsyncMock(return_value=mock_response)
            mock_client_cls.return_value.__aenter__ = AsyncMock(return_value=mock_client)
            mock_client_cls.return_value.__aexit__ = AsyncMock(return_value=False)

            with patch("asyncio.sleep", new_callable=AsyncMock):
                with pytest.raises(OCRExtractionError, match="after 3 attempts"):
                    await adapter._send_to_ocr(b"data", "test.pdf", "application/pdf")

        # Should have been called 3 times (initial + 2 retries)
        assert mock_client.post.call_count == 3

    async def test_missing_markdown_field_retries(self, adapter: OCRAdapter) -> None:
        mock_response = _make_response(200, {"text": "no markdown field"})

        with patch("httpx.AsyncClient") as mock_client_cls:
            mock_client = AsyncMock()
            mock_client.post = AsyncMock(return_value=mock_response)
            mock_client_cls.return_value.__aenter__ = AsyncMock(return_value=mock_client)
            mock_client_cls.return_value.__aexit__ = AsyncMock(return_value=False)

            with patch("asyncio.sleep", new_callable=AsyncMock):
                with pytest.raises(OCRExtractionError, match="after 3 attempts"):
                    await adapter._send_to_ocr(b"data", "test.pdf", "application/pdf")

        assert mock_client.post.call_count == 3

    async def test_connect_error_retries(self, adapter: OCRAdapter) -> None:
        with patch("httpx.AsyncClient") as mock_client_cls:
            mock_client = AsyncMock()
            mock_client.post = AsyncMock(
                side_effect=httpx.ConnectError("Connection refused")
            )
            mock_client_cls.return_value.__aenter__ = AsyncMock(return_value=mock_client)
            mock_client_cls.return_value.__aexit__ = AsyncMock(return_value=False)

            with patch("asyncio.sleep", new_callable=AsyncMock):
                with pytest.raises(OCRExtractionError, match="after 3 attempts"):
                    await adapter._send_to_ocr(b"data", "test.pdf", "application/pdf")

        assert mock_client.post.call_count == 3

    async def test_timeout_error_retries(self, adapter: OCRAdapter) -> None:
        with patch("httpx.AsyncClient") as mock_client_cls:
            mock_client = AsyncMock()
            mock_client.post = AsyncMock(
                side_effect=httpx.ReadTimeout("Read timed out")
            )
            mock_client_cls.return_value.__aenter__ = AsyncMock(return_value=mock_client)
            mock_client_cls.return_value.__aexit__ = AsyncMock(return_value=False)

            with patch("asyncio.sleep", new_callable=AsyncMock):
                with pytest.raises(OCRExtractionError, match="after 3 attempts"):
                    await adapter._send_to_ocr(b"data", "test.pdf", "application/pdf")

        assert mock_client.post.call_count == 3

    async def test_exponential_backoff_delays(self, adapter: OCRAdapter) -> None:
        mock_response = _make_response(500, {"error": "fail"})

        with patch("httpx.AsyncClient") as mock_client_cls:
            mock_client = AsyncMock()
            mock_client.post = AsyncMock(return_value=mock_response)
            mock_client_cls.return_value.__aenter__ = AsyncMock(return_value=mock_client)
            mock_client_cls.return_value.__aexit__ = AsyncMock(return_value=False)

            with patch("asyncio.sleep", new_callable=AsyncMock) as mock_sleep:
                with pytest.raises(OCRExtractionError):
                    await adapter._send_to_ocr(b"data", "test.pdf", "application/pdf")

        # Backoff: 5s * 2^0 = 5s, 5s * 2^1 = 10s (no sleep after last attempt)
        assert mock_sleep.call_count == 2
        mock_sleep.assert_any_call(5.0)
        mock_sleep.assert_any_call(10.0)

    async def test_success_on_second_attempt(self, adapter: OCRAdapter) -> None:
        fail_response = _make_response(500, {"error": "temporary"})
        success_response = _make_response(200, {"markdown": "Extracted!"})

        with patch("httpx.AsyncClient") as mock_client_cls:
            mock_client = AsyncMock()
            mock_client.post = AsyncMock(
                side_effect=[fail_response, success_response]
            )
            mock_client_cls.return_value.__aenter__ = AsyncMock(return_value=mock_client)
            mock_client_cls.return_value.__aexit__ = AsyncMock(return_value=False)

            with patch("asyncio.sleep", new_callable=AsyncMock):
                result = await adapter._send_to_ocr(b"data", "test.pdf", "application/pdf")

        assert result == "Extracted!"
        assert mock_client.post.call_count == 2


class TestHandlePDF:
    """Verify PDF handling with chunking logic."""

    async def test_small_pdf_sent_directly(self, adapter: OCRAdapter) -> None:
        """PDF with <= 20 pages should be sent directly without splitting."""
        import fitz

        # Create a small 5-page PDF
        doc = fitz.open()
        for _ in range(5):
            doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        with patch.object(adapter, "_send_to_ocr", new_callable=AsyncMock) as mock_send:
            mock_send.return_value = "small pdf text"
            result = await adapter._handle_pdf(pdf_bytes, "small.pdf")

        assert result == "small pdf text"
        mock_send.assert_called_once_with(pdf_bytes, "small.pdf", "application/pdf")

    async def test_large_pdf_split_into_chunks(self, adapter: OCRAdapter) -> None:
        """PDF with > 20 pages should be split into chunks."""
        import fitz

        # Create a 45-page PDF (should split into 3 chunks: 20, 20, 5)
        doc = fitz.open()
        for _ in range(45):
            doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        call_count = 0

        async def mock_send(data: bytes, filename: str, mime: str) -> str:
            nonlocal call_count
            call_count += 1
            return f"chunk {call_count}"

        with patch.object(adapter, "_send_to_ocr", side_effect=mock_send):
            result = await adapter._handle_pdf(pdf_bytes, "large.pdf")

        assert call_count == 3
        assert "chunk 1" in result
        assert "chunk 2" in result
        assert "chunk 3" in result
        assert "\n---PAGE BREAK---\n" in result

    async def test_exactly_20_pages_sent_directly(self, adapter: OCRAdapter) -> None:
        """PDF with exactly 20 pages should be sent directly."""
        import fitz

        doc = fitz.open()
        for _ in range(20):
            doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        with patch.object(adapter, "_send_to_ocr", new_callable=AsyncMock) as mock_send:
            mock_send.return_value = "20 pages text"
            result = await adapter._handle_pdf(pdf_bytes, "exact.pdf")

        assert result == "20 pages text"
        mock_send.assert_called_once()

    async def test_21_pages_split_into_two_chunks(self, adapter: OCRAdapter) -> None:
        """PDF with 21 pages should split into 2 chunks: 20 + 1."""
        import fitz

        doc = fitz.open()
        for _ in range(21):
            doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        call_count = 0

        async def mock_send(data: bytes, filename: str, mime: str) -> str:
            nonlocal call_count
            call_count += 1
            return f"chunk {call_count}"

        with patch.object(adapter, "_send_to_ocr", side_effect=mock_send):
            result = await adapter._handle_pdf(pdf_bytes, "split.pdf")

        assert call_count == 2
        assert result == "chunk 1\n---PAGE BREAK---\nchunk 2"


class TestHandleDOCX:
    """Verify DOCX handling with fallback logic."""

    async def test_docx_with_sufficient_text_returns_directly(
        self, adapter: OCRAdapter
    ) -> None:
        """DOCX with >= 50 chars of text should return extracted text directly."""
        from docx import Document
        import io

        # Create a DOCX with enough text
        doc = Document()
        doc.add_paragraph("This is a test paragraph with more than fifty characters of content for testing.")
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        with patch.object(adapter, "_send_to_ocr", new_callable=AsyncMock) as mock_send:
            result = await adapter.extract_text(
                docx_bytes,
                "test.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        # Should NOT call OCR since text extraction was sufficient
        mock_send.assert_not_called()
        assert "fifty characters" in result

    async def test_docx_with_short_text_and_large_file_falls_back_to_ocr(
        self, adapter: OCRAdapter
    ) -> None:
        """DOCX with < 50 chars AND > 500KB should fall back to OCR."""
        from docx import Document
        import io

        # Create a DOCX with minimal text
        doc = Document()
        doc.add_paragraph("Short")
        buffer = io.BytesIO()
        doc.save(buffer)
        # Pad to > 500KB
        small_docx = buffer.getvalue()
        large_docx = small_docx + b"\x00" * (600 * 1024)

        with patch.object(
            adapter, "_convert_docx_to_pdf", new_callable=AsyncMock
        ) as mock_convert:
            mock_convert.return_value = b"fake-pdf-bytes"
            with patch.object(
                adapter, "_send_to_ocr", new_callable=AsyncMock
            ) as mock_send:
                mock_send.return_value = "OCR result from converted PDF"

                # We need to patch Document to handle the padded bytes
                with patch("src.modules.recruitment.infrastructure.ocr_adapter.io.BytesIO") as mock_bytesio:
                    # Let the real BytesIO work for the Document creation
                    mock_bytesio.side_effect = io.BytesIO

                    result = await adapter._handle_docx(large_docx, "large.docx")

        assert result == "OCR result from converted PDF"

    async def test_docx_with_short_text_and_small_file_returns_text(
        self, adapter: OCRAdapter
    ) -> None:
        """DOCX with < 50 chars AND <= 500KB should return the short text."""
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("Hi")
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        # File is small (< 500KB), so should return text directly even if short
        with patch.object(adapter, "_send_to_ocr", new_callable=AsyncMock) as mock_send:
            result = await adapter._handle_docx(docx_bytes, "tiny.docx")

        mock_send.assert_not_called()
        assert result == "Hi"
